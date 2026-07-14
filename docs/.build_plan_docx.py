from pathlib import Path
import re
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

BASE = Path(__file__).resolve().parent
SRC = BASE / "房地产AI项目第二阶段三个月开发计划（修订版）.md"
OUT = BASE / "房地产AI项目第二阶段三个月开发计划（公文版）.docx"

RED = "8A1F1F"
INK = "20252B"
MUTED = "666666"
LIGHT = "F4F1EF"
LIGHT2 = "FAF9F8"
BORDER = "B8B2AE"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_borders(table, color=BORDER, size='6'):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = borders.find(qn(f'w:{edge}'))
        if tag is None:
            tag = OxmlElement(f'w:{edge}')
            borders.append(tag)
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), size)
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:cantSplit')
    trPr.append(el)

def set_fixed_table_geometry(table, widths_cm):
    table.autofit = False
    total_twips = sum(int(Cm(w).twips) for w in widths_cm)
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(total_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_cm:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(Cm(w).twips)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            wtw = int(Cm(widths_cm[idx]).twips)
            cell.width = Cm(widths_cm[idx])
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(wtw))
            tcW.set(qn('w:type'), 'dxa')

def set_font(run, name, size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('— ')
    set_font(run, 'Noto Serif CJK SC', 9, color=MUTED)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2])
    run2 = paragraph.add_run(' —')
    set_font(run2, 'Noto Serif CJK SC', 9, color=MUTED)

def add_bottom_border(paragraph, color=RED, size='12'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)

def add_cover(doc):
    for _ in range(5):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run('房地产 AI 项目')
    set_font(r, 'Noto Sans CJK SC', 18, True, RED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('第二阶段三个月开发计划')
    set_font(r, 'FZXiaoBiaoSong-B05S', 28, False, INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run('分阶段任务书（修订版）')
    set_font(r, 'STKaiti', 18, False, INK)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(4.0)
    p.paragraph_format.right_indent = Cm(4.0)
    add_bottom_border(p, RED, '16')
    for _ in range(8):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('计划周期：2026 年 7 月—9 月')
    set_font(r, 'Noto Serif CJK SC', 16, False, INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run('编制日期：2026 年 7 月')
    set_font(r, 'Noto Serif CJK SC', 16, False, INK)
    doc.add_page_break()

def add_rich_text(paragraph, text, font='Noto Serif CJK SC', size=16, color=INK):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part: continue
        bold = part.startswith('**') and part.endswith('**')
        val = part[2:-2] if bold else part
        run = paragraph.add_run(val)
        set_font(run, font, size, bold, color)

def add_body_paragraph(doc, text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_rich_text(p, text)
    return p

def add_label_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), LIGHT); pPr.append(shd)
    r = p.add_run(text.replace('**', ''))
    set_font(r, 'Noto Serif CJK SC', 14, True, INK)
    return p

def add_bullet(doc, text, numbered=False):
    style = 'List Number' if numbered else 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.first_line_indent = None
    add_rich_text(p, text, size=15)
    return p

def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    if cols == 4:
        widths = [3.2, 4.2, 4.2, 4.2]
    elif cols == 3 and '负责人' in rows[0][0]:
        widths = [2.7, 6.4, 6.7]
    elif cols == 3:
        widths = [2.6, 2.6, 10.6]
    else:
        widths = [15.8 / cols] * cols
    for j, val in enumerate(rows[0]):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(val.strip())
        set_font(r, 'Noto Sans CJK SC', 10.5, True, INK)
        set_cell_shading(cell, 'E8E1DD')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for i, rowvals in enumerate(rows[1:], 1):
        row = table.add_row()
        set_row_cant_split(row)
        for j, val in enumerate(rowvals):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (j == 0 or (cols == 3 and j == 1 and '负责人' not in rows[0][0])) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val.strip())
            set_font(r, 'Noto Serif CJK SC', 10.2, False, INK)
            if i % 2 == 0:
                set_cell_shading(cell, LIGHT2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_fixed_table_geometry(table, widths)
    set_table_borders(table)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    return table

def configure_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Noto Serif CJK SC'; normal.font.size = Pt(16)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(28)
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(0)

    h1 = styles['Heading 1']
    h1.font.name = 'Noto Sans CJK SC'; h1.font.size = Pt(16); h1.font.bold = True; h1.font.color.rgb = RGBColor.from_string(INK)
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    h1.paragraph_format.space_before = Pt(15); h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles['Heading 2']
    h2.font.name = 'Noto Sans CJK SC'; h2.font.size = Pt(15); h2.font.bold = True; h2.font.color.rgb = RGBColor.from_string(RED)
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    h2.paragraph_format.space_before = Pt(14); h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles['Heading 3']
    h3.font.name = 'STKaiti'; h3.font.size = Pt(15); h3.font.bold = True; h3.font.color.rgb = RGBColor.from_string(INK)
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), 'STKaiti')
    h3.paragraph_format.space_before = Pt(10); h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for styname in ['List Bullet', 'List Number']:
        st = styles[styname]
        st.font.name = 'Noto Serif CJK SC'; st.font.size = Pt(15)
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        st.paragraph_format.left_indent = Cm(1.1)
        st.paragraph_format.first_line_indent = Cm(-0.55)
        st.paragraph_format.space_after = Pt(3)

def setup_page(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.0); sec.bottom_margin = Cm(2.8)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.4)
    sec.header_distance = Cm(1.2); sec.footer_distance = Cm(1.2)
    sec.different_first_page_header_footer = True
    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('房地产 AI 项目｜第二阶段三个月开发计划')
    set_font(r, 'Noto Serif CJK SC', 9, False, MUTED)
    add_bottom_border(p, 'B9B0AA', '6')
    footer = sec.footer
    p = footer.paragraphs[0]
    add_page_number(p)

def main():
    doc = Document()
    setup_page(doc)
    configure_styles(doc)
    props = doc.core_properties
    props.title = '房地产 AI 项目第二阶段三个月开发计划（分阶段任务书）'
    props.subject = '2026年7月至9月开发计划'
    props.author = '项目开发组'
    props.keywords = '房地产AI, 开发计划, 物业费小程序, 销售App, 合同管理'
    add_cover(doc)

    lines = SRC.read_text(encoding='utf-8').splitlines()
    i = 2  # skip markdown title and subtitle
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1; continue
        if line.startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].startswith('|'):
                vals = [x.strip() for x in lines[i].strip().strip('|').split('|')]
                if not all(re.fullmatch(r'[-: ]+', x or '-') for x in vals):
                    table_rows.append(vals)
                i += 1
            if len(table_rows) >= 2:
                add_table(doc, table_rows)
            continue
        if line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='Heading 1')
            add_bottom_border(p, 'C9BFBA', '4')
        elif line.startswith('## 工作线'):
            doc.add_paragraph(line[3:], style='Heading 2')
        elif line.startswith('## '):
            doc.add_paragraph(line[3:], style='Heading 2')
        elif line.startswith('#### '):
            doc.add_paragraph(line[5:], style='Heading 3')
        elif line.startswith('- '):
            add_bullet(doc, line[2:])
        elif re.match(r'^\d+\.\s+', line):
            add_bullet(doc, re.sub(r'^\d+\.\s+', '', line), numbered=True)
        elif line.startswith('**') and line.endswith('**'):
            add_label_paragraph(doc, line)
        else:
            add_body_paragraph(doc, line)
        i += 1

    # Avoid orphaned headings and enable widow control everywhere.
    for p in doc.paragraphs:
        p.paragraph_format.widow_control = True
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    main()
